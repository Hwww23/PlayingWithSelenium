from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime, timedelta
from docx import Document
import time

def get_articles():
    # Set up the WebDriver
    driver = webdriver.Chrome()
    driver.get('https://www.straitstimes.com')
    driver.maximize_window()

    try:
        # List of keywords to search for
        search_keywords_list = [
            "cyber", "cybersecurity", "cyberattack", "cybercrime", "scam", "scammer", "hack",
            "ransomware", "data", "password", "malware", "CSA"
        ]

        articles = {}

        # Create a Word document to store the results
        doc = Document()
        doc.add_heading('Articles', level=1)
        doc.add_paragraph("")

        # Wait for the search button to be clickable and click it
        search_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CLASS_NAME, 'queryly-btn'))
        )
        search_button.click()

        time.sleep(5)

        # Loop over each search keyword
        for search_keywords in search_keywords_list:
            search_bar = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.ID, 'queryly_query'))
            )
            search_bar.send_keys(search_keywords + Keys.RETURN)

            time.sleep(2)

            # Click the "Sort by Date" button
            sort_by_date_button = driver.find_element(By.ID, 'sortbydate')
            sort_by_date_button.click()

            time.sleep(2)

            # Get the article titles and descriptions
            titles = driver.find_elements(By.CLASS_NAME, 'queryly_item_title')
            descriptions = driver.find_elements(By.CLASS_NAME, 'queryly_item_description')

            # Get the current date and determine the date range
            current_date = datetime.now()
            day_of_week = current_date.weekday()

            # Adjust start date based on the day of the week (Monday or other days)
            if day_of_week == 0:  # Monday, check from last Friday
                start_date = current_date - timedelta(days=3)
                start_date_mn = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
            else:
                start_date = current_date - timedelta(days=1)  # From yesterday
                start_date_mn = start_date.replace(hour=0, minute=0, second=0, microsecond=0)

            # Loop through each article
            for title, description in zip(titles, descriptions):
                try:
                    title_text = title.text.strip()
                    title_link = title.find_element(By.XPATH, './../../..//a').get_attribute('href')
                    
                    date_element = description.find_element(By.XPATH, 
                        './ancestor::div[contains(@class,"queryly_item_contentcontainer")]//span[@class="queryly_item_pubdate"]')
                    date_text = date_element.text.strip()
                    article_date = datetime.strptime(date_text, "%b %d, %Y")

                    # Check if the article date is within the desired date range
                    if start_date_mn <= article_date:
                        print(f"Title: {title_text}")
                        print(f"Link: {title_link}")
                        print(f"Date: {date_text}")
                        print("-" * 50)

                        articles[title_link] = {
                            'title': title_text,
                            'link': title_link,
                            'date': article_date,
                            'date_text': date_text  # Keep the original date text for later use
                        }

                except Exception as e:
                    print(f"Error processing item: {e}")

            # Clear the search bar for the next search
            search_bar.clear()

        # Sort the articles by date (descending order)
        sorted_articles = sorted(articles.values(), key=lambda x: x['date'], reverse=True)

        # Loop through the sorted articles and extract additional information
        for article in sorted_articles:
            driver.get(article['link'])

            try:
                # Extract author name, if available
                author_element = driver.find_element(By.CLASS_NAME, 'byline-name')
                author_name = author_element.text.strip()
            except:
                author_name = "No Author"

            # Check if the title mentions "CSA"
            csa_mention = 'CSA' in article['title']

            # Add information to the Word document
            if csa_mention:
                doc.add_paragraph(f"{article['title']} <BY: {author_name} | ST Online> {article['link']} Mention of CSA | Date: {article['date_text']}")
            else:
                doc.add_paragraph(f"{article['title']} <BY: {author_name} | ST Online> {article['link']} | Date: {article['date_text']}")

            doc.add_paragraph("")  # Add space between articles

        # Save the Word document
        doc.save("Articles.docx")
        print("Document saved successfully!")

    except Exception as e:
        print(f"Error: {e}")

    finally:
        # Close the browser
        driver.quit()

if __name__ == "__main__":
    get_articles()
